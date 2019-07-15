'''
ARCHES - a program developed to inventory and manage immovable cultural heritage.
Copyright (C) 2013 J. Paul Getty Trust and World Monuments Fund

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as
published by the Free Software Foundation, either version 3 of the
License, or (at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program. If not, see <http://www.gnu.org/licenses/>.
'''

import json
import collections
import couchdb
# import mammoth
import urlparse
from datetime import datetime
from datetime import timedelta
from django.db import transaction
from django.shortcuts import render
from django.db.models import Count
# from django.contrib.auth.models import User, Group
from django.contrib.gis.geos import MultiPolygon
from django.contrib.gis.geos import Polygon
from django.core.urlresolvers import reverse
from django.core.mail import EmailMultiAlternatives
from django.http import HttpRequest, HttpResponseNotFound
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.utils.translation import ugettext as _
from django.utils.decorators import method_decorator
from django.views.generic import View
from docx import Document
from arches.app.utils.betterJSONSerializer import JSONSerializer, JSONDeserializer
from arches.app.utils.response import JSONResponse
# from arches.app.utils.decorators import group_required
# from arches.app.utils.geo_utils import GeoUtils
# from arches.app.utils.couch import Couch
from arches.app.models import models
from arches.app.models.card import Card
from arches.app.models.resource import Resource
from arches.app.models.graph import Graph
from arches.app.models.system_settings import settings
from arches.app.datatypes.datatypes import DataTypeFactory
# from arches.app.views.base import BaseManagerView
# from arches.app.views.base import MapBaseManagerView
import arches.app.views.search as search
import os
import pprint


class FileTemplateView(View):

    doc = None
    tile_data = {}
    resource = None


    def get(self, request): 
        # data = JSONDeserializer().deserialize(request.body)
        datatype_factory = DataTypeFactory()
        template_id = request.GET.get('template_id')
        resourceinstance_id = request.GET.get('resourceinstance_id', None)
        self.resource = Resource.objects.get(resourceinstanceid=resourceinstance_id)
        self.resource.load_tiles()
        consultation_instance_id = None
        consultation = None
        for tile in self.resource.tiles: # self.resource is of communication model
            if 'a5901911-6d1e-11e9-8674-dca90488358a' in tile.data.keys(): # related-consultation nodegroup
                consultation_instance_id = tile.data['a5901911-6d1e-11e9-8674-dca90488358a'][0]

        template_path = self.get_template_path(template_id)
        self.doc = Document(template_path)
        new_file_name = None
        new_file_path = None

        if consultation_instance_id is not None:
            consultation = Resource.objects.get(resourceinstanceid=consultation_instance_id)
            consultation.load_tiles()
            if 'GLAAS Planning Letter A - No Progression - template.docx' in template_path:
                self.get_letter_A(consultation, datatype_factory)
            new_file_name = 'A_edited.docx'
            new_file_path = os.path.join(settings.APP_ROOT, 'uploadedfiles/docx', new_file_name)
            self.doc.save(new_file_path)
            # with open(new_file_path, "rb") as docx_file:
            #     result = mammoth.convert_to_html(docx_file)
            #     html = result.value # The generated HTML
            # with open(html_path, 'wb') as html_file:
            #     html_file.write(html)
            #     html_file.close()

        if resourceinstance_id is not None:
            return JSONResponse({'resource': self.resource, 'template_id': template_id })

        return HttpResponseNotFound()

    def get_template_path(self, template_id):
        template_path = None
        template_dict = {
            "a26c77ff-1d04-4b76-a45f-417f7ed24333":'', # Addit Cond Text
            "8c12a812-8000-4ec9-913d-c6fd516117f2":'', # Arch Rec Text
            "01dec356-e72e-40e6-b1b1-b847b9799d2f":'GLAAS Planning Letter A - No Progression - template.docx', # Letter A
            "320abc26-db82-44a6-be11-8d44aaa23365":'', # Letter A2
            "fd15c6c7-e94d-4914-8d51-a98bda6f4a7b":'', # Letter B1
            "8cc91474-11ce-47d9-b886-f0e3fc49d277":'GLAAS Planning Letter B2 - Predetermination - template.docx', # Letter B2
            "08bb630d-a27b-45bc-a13f-567b428018c5":'GLAAS Planning Letter C - Condition two stage - template.docx' # Letter C
            }
        for key, value in template_dict.items():
            if key == template_id:
                template_path = os.path.join(settings.APP_ROOT, 'docx', value)
                pprint(template_path)

        return template_path

    def get_letter_A(self, consultation, datatype_factory):
        template_dict = { # arbitrary strkey (from docx template): 'nodeid'
            'Case Officer':'36a6c511-6c49-11e9-b450-dca90488358a',
            'Completion Date': '0316def5-5675-11e9-8804-dca90488358a',
            'Proposal': 'f34ebbd4-53f3-11e9-b649-dca90488358a',
            'Log Date': '49f806e6-5674-11e9-a5b2-dca90488358a',
            'Action': '8b171540-6d1e-11e9-ac56-dca90488358a'
        }

        for tile in consultation.tiles:
            for key, value in template_dict.items():
                if value in tile.data:
                    print 'success!'
                    print (tile.data)
                    my_node = models.Node.objects.get(nodeid=value)
                    datatype = datatype_factory.get_instance(my_node.datatype)
                    lookup_val = datatype.get_display_value(tile, my_node)
                    self.replace_string(self.doc, key, lookup_val)

    
    def replace_string(self, document, key, v):
        if v is not None and key is not None:
            k = "{{"+key+"}}"
            doc = document
            t_style = None
            p_style = None
            run_style = None

            if len(doc.paragraphs) > 0:
                for p in doc.paragraphs:
                    if k in p.text:
                        # print (k,'key is in p:',p.text)
                        p_style = p.style
                        run_style = p.runs[0].style
                        p.text = p.text.replace(k, v)
                        p.style = p_style
                        p.runs[0].style = run_style

            if len(doc.tables) > 0:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if k in cell.text:
                                # print (k, 'key is in cell:',cell.text)
                                t_style = table.style
                                p_style = cell.paragraphs[0].style
                                run_style = cell.paragraphs[0].runs[0].style
                                cell.text = cell.text.replace(k, v)
                                table.style = t_style
                                cell.paragraphs[0].style = p_style
                                cell.paragraphs[0].runs[0].style = run_style
            
            if len(doc.sections) > 0:
                for section in doc.sections:
                    for p in section.footer.paragraphs:
                        if k in p.text:
                            p_style = p.style
                            run_style = p.runs[0].style
                            p.text = p.text.replace(k, v)
                            p.style = p_style
                            p.runs[0].style = run_style
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if k in cell.text:
                                    t_style = table.style
                                    p_style = cell.paragraphs[0].style
                                    run_style = cell.paragraphs[0].runs[0].style
                                    cell.text = cell.text.replace(k, v)
                                    table.style = t_style
                                    cell.paragraphs[0].style = p_style
                                    cell.paragraphs[0].runs[0].style = run_style
                    for p in section.header.paragraphs:
                        if k in p.text:
                            p_style = p.style
                            run_style = p.runs[0].style
                            p.text = p.text.replace(k, v)
                            p.style = p_style
                            p.runs[0].style = run_style
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if k in cell.text:
                                    t_style = table.style
                                    p_style = cell.paragraphs[0].style
                                    run_style = cell.paragraphs[0].runs[0].style
                                    cell.text = cell.text.replace(k, v)
                                    table.style = t_style
                                    cell.paragraphs[0].style = p_style
                                    cell.paragraphs[0].runs[0].style = run_style

    
    def insert_image(self, document, k, v, image_path=None, config=None):
        # going to need to write custom logic depending on how images should be placed/styled

        return True


    def insert_custom(self, document, k, v, config=None):
        # perhaps replaces {{custom_object}} with pre-determined text structure with custom style/format

        return True

